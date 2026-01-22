import subprocess
import os
import sys
import shutil

# --- DEPENDENCY CHECK ---
try:
    from docx import Document
except ImportError:
    print("\n[ERROR] 'python-docx' library not found.")
    print("Please run: pip install python-docx")
    sys.exit(1)


def extract_audio(video_path, audio_path):
    """Uses FFmpeg to extract audio from the video file."""
    print(f"\n--- 1. EXTRACTING AUDIO ---")
    print(f"Input Video: {os.path.basename(video_path)}")
    
    # FFmpeg arguments:
    # -y: Always overwrite output (prevents hanging)
    # -vn: Disable video recording
    # -acodec libmp3lame: Standard MP3 codec
    # -q:a 2: High quality audio
    command = [
        'ffmpeg', '-y', '-i', video_path, '-vn',
        '-acodec', 'libmp3lame', '-q:a', '2', audio_path
    ]
    
    try:
        subprocess.run(command, check=True, capture_output=True)
        print("Status: Audio extraction successful.")
    except subprocess.CalledProcessError as e:
        print(f"Error: FFmpeg failed. {e}")
        sys.exit(1)

def run_whisper_ai(audio_path, language_code, do_translate):
    """Loads Whisper and processes the audio for transcription or translation."""
    task = "translate" if do_translate else "transcribe"
    label = "Translation (to English)" if do_translate else f"Transcription ({language_code})"
    
    print(f"\n--- 2. RUNNING WHISPER AI ---")
    print(f"Task: {label}")
    print("Note: The AI model may take a minute to process...")
    
    try:
        import whisper
        # 'base' model is the perfect balance for Mac performance
        model = whisper.load_model("base")
        
        # We pass the language and the task type
        result = model.transcribe(audio_path, language=language_code, task=task)
        return result["text"]
    except ImportError:
        print("\n[ERROR] 'openai-whisper' not found. Run: pip install openai-whisper")
        sys.exit(1)
    except Exception as e:
        print(f"Error during AI processing: {e}")
        sys.exit(1)

def save_to_word(text, output_path):
    """Saves the final text result into a Microsoft Word document."""
    print(f"\n--- 3. SAVING WORD DOCUMENT ---")
    try:
        doc = Document()
        doc.add_heading('Whisper AI Transcription Result', 0)
        doc.add_paragraph(text)
        doc.save(output_path)
        print(f"Status: Saved successfully to {os.path.basename(output_path)}")
    except Exception as e:
        print(f"Error saving Word doc: {e}")

def main():
    # Verify we have the right number of arguments
    # Usage: python3 transcribe_video.py <file> <lang> [translate]
    if len(sys.argv) < 3:
        print("\n" + "="*40)
        print("VIDEO TRANSCRIPTION PIPELINE")
        print("="*40)
        print("Usage:  python3 transcribe_video.py <video_path> <lang_code> [translate]")
        print("\nExamples:")
        print("  Just Transcribe:  python3 transcribe_video.py video.mp4 ja")
        print("  Translate to EN:  python3 transcribe_video.py video.mp4 ja translate")
        print("="*40 + "\n")
        return


    # Get arguments and normalize paths
    video_input = os.path.abspath(sys.argv[1])
    source_lang = sys.argv[2]
    wants_translation = len(sys.argv) > 3 and sys.argv[3].lower() == "translate"

    # Set up file names
    file_root = os.path.splitext(video_input)[0]
    temp_audio = f"{file_root}_temp_audio.mp3"
    
    if wants_translation:
        final_doc = f"{file_root}_translated.docx"
    else:
        final_doc = f"{file_root}_transcript_{source_lang}.docx"

    if not os.path.exists(video_input):
        print(f"File Not Found: {video_input}")
        return

    # Start the engine
    try:
        extract_audio(video_input, temp_audio)
        ai_text = run_whisper_ai(temp_audio, source_lang, wants_translation)
        save_to_word(ai_text, final_doc)
        
        # Clean up the temporary audio file
        if os.path.exists(temp_audio):
            os.remove(temp_audio)
            
        print("\n" + "*"*30)
        print("  PIPELINE FINISHED SUCCESSFULLY")
        print("*"*30 + "\n")

    except KeyboardInterrupt:
        print("\n\nProcess stopped by user.")
        if os.path.exists(temp_audio):
            os.remove(temp_audio)

if __name__ == "__main__":
    main()