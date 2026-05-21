"""
DOCX document builder for AI Transcription Reports.

Document layout:
  ┌──────────────────────────────────────────────────┐
  │  AI Transcription Report          (Title style)  │
  │  source_filename.mp4              (14pt gray)     │
  │  ──────────────────────────────────────────────  │
  │  [borderless 2-column metadata table]            │
  │    Generated        2026-05-21  14:30:00         │
  │    Mode             Transcription                │
  │    Source Language  en                           │
  │    Target Language  English   ← only if xlated  │
  │    Whisper Model    base                         │
  │    Output File      name_transcript_en.docx      │
  │  ──────────────────────────────────────────────  │
  ├──────────────────────────────────────────────────┤
  │  [page break]                                    │
  ├──────────────────────────────────────────────────┤
  │  Transcript                       (Heading 1)    │
  │                                                  │
  │  With --timestamps:                              │
  │    [00:00:04 – 00:00:12]  (9pt bold light-gray) │
  │    Segment text paragraph. (11pt Calibri)        │
  │                                                  │
  │  Without --timestamps:                           │
  │    Paragraph 1 (up to 4 sentences)              │
  │    Paragraph 2 …                                │
  └──────────────────────────────────────────────────┘

Only python-docx is required. XML helpers are used solely to remove
table borders — a standard python-docx pattern with no extra deps.
"""

import os
import re
from datetime import datetime

from app.errors import DependencyError, ExportError
from app.formatter import seconds_to_hms


# ── Colour palette ────────────────────────────────────────────────────────────

_BODY_FONT    = "Calibri"
_HEADING_FONT = "Calibri Light"
_CLR_MID      = "6B7280"   # gray  – metadata labels
_CLR_LIGHT    = "9CA3AF"   # light gray – timestamp labels, separator lines


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _rgb(hex_str):
    """Convert a 6-char hex string like '6B7280' to an RGBColor object."""
    from docx.shared import RGBColor
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _spacing(para, before=0, after=6):
    """Set space-before / space-after in points on a paragraph."""
    from docx.shared import Pt
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _add_rule(doc):
    """Insert an empty paragraph with a thin light-gray bottom border (visual rule)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _spacing(p, before=2, after=2)


# ── Cover section helpers ─────────────────────────────────────────────────────

def _add_metadata_table(doc, rows):
    """Borderless 2-column table for key / value metadata rows on the cover page."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt

    table = doc.add_table(rows=len(rows), cols=2)

    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells

        key_run = cells[0].paragraphs[0].add_run(label)
        key_run.bold = True
        key_run.font.name = _BODY_FONT
        key_run.font.size = Pt(10)
        key_run.font.color.rgb = _rgb(_CLR_MID)
        cells[0].paragraphs[0].paragraph_format.space_after = Pt(3)

        val_run = cells[1].paragraphs[0].add_run(value)
        val_run.font.name = _BODY_FONT
        val_run.font.size = Pt(10)
        cells[1].paragraphs[0].paragraph_format.space_after = Pt(3)

    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.2)

    # Strip all visible borders from the table.
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _add_cover(doc, metadata, output_path):
    """Write the title, subtitle, and metadata table onto the cover page."""
    from docx.shared import Pt

    title = doc.add_heading("AI Transcription Report", 0)
    _spacing(title, before=0, after=4)

    source_file = (metadata or {}).get("source_file", "")
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(source_file)
    sub_run.font.name  = _HEADING_FONT
    sub_run.font.size  = Pt(14)
    sub_run.font.color.rgb = _rgb(_CLR_MID)
    _spacing(subtitle, before=0, after=10)

    _add_rule(doc)

    if metadata:
        is_translation = metadata.get("translated", False)
        lang  = metadata.get("language", "auto")
        model = metadata.get("model",    "base")
        generated = datetime.now().strftime("%Y-%m-%d   %H:%M:%S")
        mode = "Translation to English" if is_translation else "Transcription"

        rows = [
            ("Generated",       generated),
            ("Mode",            mode),
            ("Source Language", lang),
        ]
        if is_translation:
            rows.append(("Target Language", "English"))
        rows.extend([
            ("Whisper Model", model),
            ("Output File",   os.path.basename(output_path)),
        ])
        cleaning_steps = metadata.get("cleaning_steps", [])
        if cleaning_steps:
            rows.append(("Cleaning", ", ".join(cleaning_steps)))

        doc.add_paragraph()          # breathing room above table
        _add_metadata_table(doc, rows)

    _add_rule(doc)


# ── Transcript section helpers ────────────────────────────────────────────────

def _split_into_paragraphs(text, sentences_per_para=4):
    """Group a text blob into readable chunks by splitting at sentence boundaries.

    Falls back to returning the whole text if no sentence-end / capital-start
    boundaries are found (e.g. non-Latin scripts or very short text).
    """
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", text.strip())
    if len(sentences) <= 1:
        return [text.strip()]
    paras = []
    for i in range(0, len(sentences), sentences_per_para):
        chunk = " ".join(s.strip() for s in sentences[i:i + sentences_per_para] if s.strip())
        if chunk:
            paras.append(chunk)
    return paras or [text.strip()]


def _add_timestamped_transcript(doc, segments):
    """One bold-gray timestamp label + one body paragraph per Whisper segment."""
    from docx.shared import Pt

    for seg in segments:
        seg_text = seg["text"].strip()
        if not seg_text:
            continue

        start = seconds_to_hms(seg["start"])
        end   = seconds_to_hms(seg["end"])

        # Timestamp label – bold, light gray, small
        ts_para = doc.add_paragraph()
        ts_run  = ts_para.add_run(f"[{start} – {end}]")   # en-dash
        ts_run.bold = True
        ts_run.font.name = _BODY_FONT
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = _rgb(_CLR_LIGHT)
        _spacing(ts_para, before=12, after=1)

        # Segment body text
        body = doc.add_paragraph(seg_text)
        for run in body.runs:
            run.font.name = _BODY_FONT
            run.font.size = Pt(11)
        _spacing(body, before=0, after=6)


def _add_plain_transcript(doc, text):
    """Paragraph-per-chunk body text when timestamps are disabled.

    If the text already contains double-newline paragraph breaks (inserted
    by --paragraphs upstream), those boundaries are honoured exactly.
    Otherwise sentences are grouped internally by _split_into_paragraphs.
    """
    from docx.shared import Pt

    if "\n\n" in text:
        chunks = [c.strip() for c in text.split("\n\n") if c.strip()]
    else:
        chunks = _split_into_paragraphs(text)

    for chunk in chunks:
        para = doc.add_paragraph(chunk)
        for run in para.runs:
            run.font.name = _BODY_FONT
            run.font.size = Pt(11)
        _spacing(para, before=0, after=10)


def _add_transcript(doc, text, segments):
    """'Transcript' heading followed by timestamped or plain body content."""
    from docx.shared import Pt

    heading = doc.add_heading("Transcript", level=1)
    _spacing(heading, before=0, after=8)

    if segments:
        _add_timestamped_transcript(doc, segments)
    else:
        _add_plain_transcript(doc, text)


# ── Notes / summary section ───────────────────────────────────────────────────

def _add_notes(doc, notes):
    """Render summary and notes sections before the transcript.

    Iterates over notes["layout"] — a list of (data_key, label, render_type) —
    and renders each non-empty section as a Heading 1 block followed by
    its content.  Empty sections are silently skipped.
    """
    from docx.shared import Pt

    for data_key, label, render_type in notes.get("layout", []):
        content = notes.get(data_key)
        if not content:
            continue  # skip empty sections gracefully

        heading = doc.add_heading(label, level=1)
        _spacing(heading, before=12, after=6)

        if render_type == "text":
            # Prose paragraph (used for summary)
            para = doc.add_paragraph(content)
            for run in para.runs:
                run.font.name = _BODY_FONT
                run.font.size = Pt(11)
            _spacing(para, before=0, after=10)

        elif render_type == "list":
            # Bulleted list of sentences
            for item in content:
                if not item:
                    continue
                para = doc.add_paragraph()
                run = para.add_run(f"•  {item}")   # bullet character
                run.font.name = _BODY_FONT
                run.font.size = Pt(11)
                _spacing(para, before=0, after=4)

        elif render_type == "keyword_list":
            # Short terms as a comma-separated paragraph
            if isinstance(content, list):
                display = ",  ".join(content)
            else:
                display = str(content)
            para = doc.add_paragraph(display)
            for run in para.runs:
                run.font.name  = _BODY_FONT
                run.font.size  = Pt(10)
                run.font.color.rgb = _rgb(_CLR_MID)
            _spacing(para, before=0, after=10)

    # Thin rule before the transcript section
    _add_rule(doc)


# ── Public entry point ────────────────────────────────────────────────────────

def build_docx(text, output_path, segments, metadata):
    """Assemble and save the complete AI Transcription Report .docx file.

    Args:
        text:        Full plain-text transcript (used for plain and as JSON fallback).
        output_path: Absolute path to write the .docx file.
        segments:    Whisper segment list (passed when --timestamps is active),
                     or None for plain paragraph output.
        metadata:    Dict with keys: source_file, language, model, translated,
                     has_timestamps.  May be None for old-style calls.
                     If metadata["notes"] is present the notes sections are
                     rendered before the transcript.
    """
    try:
        from docx import Document
    except ImportError:
        raise DependencyError(
            "The 'python-docx' package is not installed.\n"
            "Activate your virtual environment, then run:\n"
            "  pip install python-docx"
        )

    try:
        notes = (metadata or {}).get("notes")
        doc = Document()
        _add_cover(doc, metadata, output_path)
        doc.add_page_break()
        if notes:
            _add_notes(doc, notes)
        _add_transcript(doc, text, segments)
        doc.save(output_path)
    except (DependencyError, ExportError):
        raise
    except Exception as e:
        raise ExportError(
            f"Could not build Word document: {os.path.basename(output_path)}\n"
            "Run with --debug for details.",
            details=str(e),
        ) from e
